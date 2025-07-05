import pdfplumber
import docx

def extract_text(file_input):
    # If input is a string path
    if isinstance(file_input, str):
        if file_input.lower().endswith('.pdf'):
            with pdfplumber.open(file_input) as pdf:
                text = ''
                for page in pdf.pages:
                    text += page.extract_text() or ''
            return text

        elif file_input.lower().endswith('.docx'):
            doc = docx.Document(file_input)
            return '\n'.join([p.text for p in doc.paragraphs])

    # If input is an uploaded file (Streamlit)
    else:
        file_name = file_input.name
        if file_name.endswith('.pdf'):
            with pdfplumber.open(file_input) as pdf:
                text = ''
                for page in pdf.pages:
                    text += page.extract_text() or ''
            return text

        elif file_name.endswith('.docx'):
            doc = docx.Document(file_input)
            return '\n'.join([p.text for p in doc.paragraphs])

    raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
